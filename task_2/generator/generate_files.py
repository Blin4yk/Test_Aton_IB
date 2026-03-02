import os
import shutil
import subprocess
import zipfile

from docx import Document
from faker import Faker
from openpyxl import Workbook as XlsxWorkbook
from reportlab.pdfgen import canvas
from xlwt import Workbook as XlsWorkbook

from logger import logger

fake = Faker('ru_RU')


class FileGenerator:
    def __init__(self, output_dir="./test_files"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def random_text(self, paragraphs=3):
        return "\n\n".join(fake.paragraphs(paragraphs))

    def generate_doc(self, path):
        # Для .doc создаём простой текстовый файл (эмуляция)
        with open(path, 'w', encoding='cp1251') as f:
            f.write(self.random_text())

    def generate_docx(self, path):
        doc = Document()
        doc.add_heading(fake.sentence(), 0)
        for _ in range(3):
            doc.add_paragraph(fake.paragraph())
        doc.save(path)

    def generate_xls(self, path):
        wb = XlsWorkbook()
        sheet = wb.add_sheet('Sheet1')
        for i in range(5):
            for j in range(3):
                sheet.write(i, j, fake.word())
        wb.save(path)

    def generate_xlsx(self, path):
        wb = XlsxWorkbook()
        ws = wb.active
        for i in range(5):
            for j in range(3):
                ws.cell(row=i + 1, column=j + 1, value=fake.word())
        wb.save(path)

    def generate_pdf(self, path):
        c = canvas.Canvas(path)
        fake = Faker('en_US') # Заменили на английский, тк на устройстве не стоит предустановленных шрифтов
        c.drawString(100, 750, fake.sentence())
        y = 700
        for _ in range(5):
            c.drawString(100, y, fake.paragraph())
            y -= 50
        c.save()

    def create_archive(self, archive_path, files_to_archive, archive_format='zip'):
        if archive_format == 'zip':
            with zipfile.ZipFile(archive_path, 'w') as zf:
                for file in files_to_archive:
                    zf.write(file, os.path.basename(file))
        elif archive_format == 'rar':
            # Требуется установленный rar
            subprocess.run(['rar', 'a', archive_path] + files_to_archive, check=False)
        elif archive_format == '7z':
            subprocess.run(['7z', 'a', archive_path] + files_to_archive, check=False)

    def generate_all(self):
        # Создаём отдельные файлы
        files = []
        doc_file = os.path.join(self.output_dir, "test1.doc")
        self.generate_doc(doc_file)
        files.append(doc_file)

        docx_file = os.path.join(self.output_dir, "test2.docx")
        self.generate_docx(docx_file)
        files.append(docx_file)

        xls_file = os.path.join(self.output_dir, "test3.xls")
        self.generate_xls(xls_file)
        files.append(xls_file)

        xlsx_file = os.path.join(self.output_dir, "test4.xlsx")
        self.generate_xlsx(xlsx_file)
        files.append(xlsx_file)

        pdf_file = os.path.join(self.output_dir, "test5.pdf")
        self.generate_pdf(pdf_file)
        files.append(pdf_file)

        # Поддиректория с файлами
        subdir = os.path.join(self.output_dir, "subdir")
        os.makedirs(subdir, exist_ok=True)
        sub_doc = os.path.join(subdir, "sub_test.doc")
        self.generate_doc(sub_doc)
        sub_pdf = os.path.join(subdir, "sub_test.pdf")
        self.generate_pdf(sub_pdf)

        # Архивы
        archive_files = [doc_file, docx_file, xls_file]
        zip_path = os.path.join(self.output_dir, "archive.zip")
        self.create_archive(zip_path, archive_files, 'zip')

        # Вложенный архив: архив внутри архива
        nested_zip_path = os.path.join(self.output_dir, "nested.zip")
        with zipfile.ZipFile(nested_zip_path, 'w') as zf:
            zf.write(zip_path, os.path.basename(zip_path))

        # Если есть rar/7z утилиты, можно создать дополнительные архивы
        if shutil.which('rar'):
            rar_path = os.path.join(self.output_dir, "archive.rar")
            self.create_archive(rar_path, archive_files, 'rar')
        if shutil.which('7z'):
            sevenz_path = os.path.join(self.output_dir, "archive.7z")
            self.create_archive(sevenz_path, archive_files, '7z')

        logger.info("Генерация завершена.")
